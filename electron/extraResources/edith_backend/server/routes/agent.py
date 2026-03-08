"""
§AGENT: Agent Routes — Intent, Autonomous Execution, Profile
=============================================================
The agentic layer that makes E.D.I.T.H. feel like an AI assistant
instead of a collection of buttons.
"""
from __future__ import annotations

import asyncio
import json
import logging
import time

from fastapi import APIRouter, Request
from fastapi.responses import JSONResponse, StreamingResponse

log = logging.getLogger("edith.routes.agent")
router = APIRouter()


def _error(status: int, code: str, detail: str):
    return JSONResponse(status_code=status, content={"error": code, "detail": detail})


# ═══════════════════════════════════════════════════════════════════
# §1: INTENT CLASSIFICATION — "What does the user want?"
# ═══════════════════════════════════════════════════════════════════

@router.post("/api/intent/classify", tags=["Agent"])
async def classify_intent(request: Request):
    """Classify user intent without executing.

    Returns the detected intent, confidence, and planned pipeline steps.
    Useful for UI: show the user what Winnie plans to do before doing it.
    """
    body = await request.json()
    query = body.get("query", body.get("text", ""))
    context = body.get("context", "")

    if not query:
        return _error(422, "no_query", "Provide a 'query' to classify")

    from server.intent_router import route_intent
    result = route_intent(query, context=context)

    # Record to profile
    try:
        from server.research_profile import get_profile
        profile = get_profile()
        profile.record_query(query, intent=result["intent"])
        # Extract topic from query
        topic = _extract_topic(query)
        if topic:
            profile.record_topic(topic)
    except Exception:
        pass

    return result


# ═══════════════════════════════════════════════════════════════════
# §2: AUTO-EXECUTE — "Just do it"
# Classify intent + execute the pipeline in one call
# ═══════════════════════════════════════════════════════════════════

@router.post("/api/intent/run", tags=["Agent"])
async def intent_run(request: Request):
    """Classify intent AND execute the selected pipeline.

    This is the "just tell Winnie what you want" endpoint.
    One call: NL request → intent classification → pipeline execution → result.
    """
    body = await request.json()
    query = body.get("query", body.get("text", ""))
    session_id = body.get("session_id")
    extra_params = body.get("params", {})

    if not query:
        return _error(422, "no_query", "Provide a 'query'")

    t0 = time.time()

    # Step 1: Classify intent
    from server.intent_router import route_intent
    intent_result = route_intent(query)
    pipeline_steps = intent_result.get("pipeline", [])

    # If no pipeline (chat intent), return classification only
    if not pipeline_steps:
        return {
            "intent": intent_result,
            "action": "chat",
            "message": "This is a general question — use the chat endpoint.",
            "elapsed_s": round(time.time() - t0, 2),
        }

    # Step 2: Execute pipeline
    from server.routes.pipeline import _execute_step, _get_registry
    registry = _get_registry()

    # Build context from query + extra params + research profile
    context = {"text": query, **extra_params}

    # Inject research profile context
    try:
        from server.research_profile import get_profile
        profile = get_profile()
        profile_ctx = profile.get_context_string()
        if profile_ctx:
            context["_profile_context"] = profile_ctx
        lang = profile.get_profile().get("preferred_language", "stata")
        context.setdefault("language", lang)
    except Exception:
        pass

    # Load session context
    if session_id:
        try:
            from server.agentic import sessions
            prev = sessions.get(session_id)
            if prev:
                last = prev[-1] if prev else {}
                for k, v in last.items():
                    if k not in context and not k.startswith("_"):
                        context[k] = v
        except ImportError:
            pass

    # Execute steps with error recovery
    results = []
    for i, step_name in enumerate(pipeline_steps):
        if step_name not in registry:
            results.append({"step": step_name, "error": f"Unknown step: {step_name}"})
            continue

        step_t0 = time.time()
        try:
            result = await asyncio.to_thread(_execute_step, step_name, context)
        except Exception as e:
            result = {"error": str(e)[:200]}

        elapsed = round(time.time() - step_t0, 2)

        step_record = {
            "step": step_name,
            "label": registry.get(step_name, {}).get("label", step_name),
            "order": i + 1,
            "elapsed_s": elapsed,
            "success": "error" not in result,
        }

        if result.get("error"):
            # Error recovery: try to continue past failed steps
            step_record["error"] = result["error"]
            step_record["recovered"] = True
            results.append(step_record)
            log.warning(f"Pipeline step '{step_name}' failed, continuing: {result['error'][:100]}")
            continue

        step_record["result_keys"] = list(result.keys())
        results.append(step_record)

        # Merge result into context for next step
        for k, v in result.items():
            if not k.startswith("_"):
                context[k] = v

    total_elapsed = round(time.time() - t0, 2)

    # Record tool usage
    try:
        from server.research_profile import get_profile
        profile = get_profile()
        for step in pipeline_steps:
            profile.record_tool_use(step)
    except Exception:
        pass

    # Save to session
    if session_id:
        try:
            from server.agentic import sessions
            sessions.append(session_id, {
                "role": "agent",
                "content": f"Ran: {query[:100]} → {' → '.join(pipeline_steps)}",
                "intent": intent_result["intent"],
            })
        except ImportError:
            pass

    return {
        "intent": intent_result,
        "pipeline": pipeline_steps,
        "steps_completed": len([r for r in results if r.get("success")]),
        "steps_total": len(pipeline_steps),
        "total_elapsed_s": total_elapsed,
        "results": results,
        "final_context": _safe_context(context),
        "session_id": session_id,
    }


def _safe_context(ctx: dict) -> dict:
    """Serialize context dict safely — drop non-serializable values."""
    safe = {}
    for k, v in ctx.items():
        if k.startswith("_"):
            continue
        if isinstance(v, (str, int, float, bool, type(None))):
            safe[k] = v
        elif isinstance(v, (list, dict)):
            try:
                import json
                json.dumps(v)  # test serializable
                safe[k] = v
            except (TypeError, ValueError):
                safe[k] = str(v)[:500]
        else:
            safe[k] = str(v)[:500]
    return safe



# ═══════════════════════════════════════════════════════════════════
# §4: AUTONOMOUS PLANNER — "Winnie plans the steps herself"
# ═══════════════════════════════════════════════════════════════════

@router.post("/api/agent/plan", tags=["Agent"])
async def agent_plan(request: Request):
    """Let the LLM plan a custom pipeline from a goal.

    Unlike intent/classify (which maps to predefined pipelines),
    this asks the LLM to pick which steps to run and in what order.
    """
    body = await request.json()
    goal = body.get("goal", body.get("query", ""))
    constraints = body.get("constraints", {})

    if not goal:
        return _error(422, "no_goal", "Provide a 'goal'")

    from server.intent_router import plan_execution
    plan = await asyncio.to_thread(plan_execution, goal, constraints)
    return plan


@router.post("/api/agent/run", tags=["Agent"])
async def agent_run(request: Request):
    """Autonomous agent: plan + execute in one call.

    1. LLM plans the steps
    2. Pipeline executes them
    3. Returns unified result

    This is the most "agentic" endpoint — the user provides a goal,
    and Winnie figures out everything else.
    """
    body = await request.json()
    goal = body.get("goal", body.get("query", ""))
    constraints = body.get("constraints", {})
    session_id = body.get("session_id")

    if not goal:
        return _error(422, "no_goal", "Provide a 'goal'")

    t0 = time.time()

    # Step 1: Plan
    from server.intent_router import plan_execution
    plan = await asyncio.to_thread(plan_execution, goal, constraints)
    steps = plan.get("steps", [])

    if not steps:
        return {
            "plan": plan,
            "action": "no_steps",
            "message": "Could not determine steps for this goal.",
            "elapsed_s": round(time.time() - t0, 2),
        }

    # Step 2: Execute via pipeline
    from server.routes.pipeline import _execute_step, _get_registry
    registry = _get_registry()

    context = {"text": goal, **constraints}

    # Inject profile context
    try:
        from server.research_profile import get_profile
        profile = get_profile()
        context.setdefault("language", profile.get_profile().get("preferred_language", "stata"))
    except Exception:
        pass

    results = []
    for i, step_name in enumerate(steps):
        if step_name not in registry:
            results.append({"step": step_name, "error": "Unknown", "recovered": True})
            continue

        step_t0 = time.time()
        try:
            result = await asyncio.to_thread(_execute_step, step_name, context)
        except Exception as e:
            result = {"error": str(e)[:200]}

        elapsed = round(time.time() - step_t0, 2)
        step_ok = "error" not in result

        results.append({
            "step": step_name,
            "label": registry.get(step_name, {}).get("label", step_name),
            "order": i + 1,
            "elapsed_s": elapsed,
            "success": step_ok,
            **({"error": result["error"], "recovered": True} if not step_ok else {}),
        })

        if step_ok:
            for k, v in result.items():
                if not k.startswith("_"):
                    context[k] = v

    total_elapsed = round(time.time() - t0, 2)

    # Record
    try:
        from server.research_profile import get_profile
        profile = get_profile()
        profile.record_query(goal, intent="autonomous_agent")
        for s in steps:
            profile.record_tool_use(s)
    except Exception:
        pass

    return {
        "plan": plan,
        "steps_completed": len([r for r in results if r.get("success")]),
        "steps_total": len(steps),
        "total_elapsed_s": total_elapsed,
        "results": results,
        "final_context": _safe_context(context),
        "session_id": session_id,
    }


# ═══════════════════════════════════════════════════════════════════
# §5: RESEARCH PROFILE — Read / Update
# ═══════════════════════════════════════════════════════════════════

@router.get("/api/profile", tags=["Agent"])
async def get_research_profile():
    """Get the researcher's profile."""
    try:
        from server.research_profile import get_profile
        return get_profile().get_profile()
    except Exception as e:
        return {"error": str(e)}


@router.post("/api/profile/update", tags=["Agent"])
async def update_research_profile(request: Request):
    """Update profile fields (dissertation_topic, preferred_language, etc.)."""
    body = await request.json()
    try:
        from server.research_profile import get_profile
        updated = get_profile().update_profile(body)
        return {"updated": True, "profile": updated}
    except Exception as e:
        return _error(500, "profile_error", str(e))


@router.get("/api/profile/context", tags=["Agent"])
async def get_profile_context():
    """Get the profile context string (what gets injected into LLM prompts)."""
    try:
        from server.research_profile import get_profile
        return {"context": get_profile().get_context_string()}
    except Exception as e:
        return {"context": "", "error": str(e)}


@router.post("/api/profile/feedback", tags=["Agent"])
async def add_profile_feedback(request: Request):
    """Record committee/advisor feedback."""
    body = await request.json()
    feedback = body.get("feedback", "")
    source = body.get("source", "advisor")

    if not feedback:
        return _error(422, "no_feedback", "Provide 'feedback' text")

    try:
        from server.research_profile import get_profile
        get_profile().add_feedback(feedback, source)
        return {"recorded": True}
    except Exception as e:
        return _error(500, "feedback_error", str(e))


# ═══════════════════════════════════════════════════════════════════
# §6: REACTIVE CONTEXT — Accept metadata about user's current state
# ═══════════════════════════════════════════════════════════════════

@router.post("/api/agent/context", tags=["Agent"])
async def update_reactive_context(request: Request):
    """Accept context about what the user is currently doing.

    The frontend can call this to tell Winnie:
    - What file is open
    - What panel is active
    - What time of day it is
    - What the user just did

    Winnie stores this and uses it to personalize responses.
    """
    body = await request.json()

    # Store in session
    session_id = body.get("session_id", "reactive_default")
    try:
        from server.agentic import sessions
        sessions.append(session_id, {
            "role": "context",
            "content": json.dumps({
                "active_panel": body.get("active_panel"),
                "active_file": body.get("active_file"),
                "recent_action": body.get("recent_action"),
                "time_of_day": body.get("time_of_day"),
            }),
        })
        return {"stored": True, "session_id": session_id}
    except Exception as e:
        return {"stored": False, "error": str(e)}


# ── Topic extraction helper ──────────────────────────────────────

def _extract_topic(query: str) -> str:
    """Extract the primary topic from a query."""
    # Remove common query prefixes
    q = query.lower().strip()
    for prefix in ["analyze", "explain", "what is", "how does", "why does",
                    "compare", "evaluate", "discuss", "describe", "can you"]:
        if q.startswith(prefix):
            q = q[len(prefix):].strip()
    # Take first meaningful phrase
    words = q.split()
    if len(words) <= 5:
        return " ".join(words)
    return " ".join(words[:5])


# ═══════════════════════════════════════════════════════════════════
# §8: BACKGROUND TASK QUEUE — "Run overnight"
# ═══════════════════════════════════════════════════════════════════

_task_queue: list[dict] = []
_task_counter = 0


async def _run_background_task(task_id: str, query: str) -> None:
    """Run a pipeline in the background and store the result."""
    global _task_queue
    for t in _task_queue:
        if t["id"] == task_id:
            t["status"] = "running"
            break
    _save_queue(_task_queue)

    try:
        from server.intent_router import route_intent
        from server.routes.pipeline import _execute_step, _get_registry

        intent = route_intent(query)
        pipeline = intent.get("pipeline", [])
        registry = _get_registry()
        context = {"text": query}

        for step_name in pipeline:
            if step_name not in registry:
                continue
            try:
                result = await asyncio.to_thread(_execute_step, step_name, context)
                for k, v in result.items():
                    if not k.startswith("_"):
                        context[k] = v
            except Exception as e:
                log.warning(f"Background step '{step_name}' failed: {e}")

        for t in _task_queue:
            if t["id"] == task_id:
                t["status"] = "done"
                t["result"] = _safe_context(context)
                break
        _save_queue(_task_queue)

    except Exception as e:
        for t in _task_queue:
            if t["id"] == task_id:
                t["status"] = "error"
                t["error"] = str(e)[:200]
                break
        _save_queue(_task_queue)


@router.post("/api/agent/queue", tags=["Agent"])
async def queue_task(request: Request):
    """Queue a task for background execution."""
    global _task_counter
    body = await request.json()
    task_text = body.get("task", body.get("query", ""))
    if not task_text:
        return _error(422, "no_task", "Provide a 'task'")

    _task_counter += 1
    task_id = f"bg-{_task_counter}"

    task_entry = {
        "id": task_id,
        "task": task_text,
        "status": "queued",
        "queued_at": time.strftime("%Y-%m-%dT%H:%M:%SZ"),
    }
    _task_queue.insert(0, task_entry)
    _save_queue(_task_queue)

    # Launch in background
    asyncio.create_task(_run_background_task(task_id, task_text))

    return {"id": task_id, "status": "queued"}


@router.get("/api/agent/queue/status", tags=["Agent"])
async def queue_status():
    """Return all queued/running/completed tasks."""
    return {"tasks": _task_queue[:20]}


@router.get("/api/agent/briefing", tags=["Agent"])
async def get_briefing():
    """Return the latest overnight briefing."""
    try:
        from server.overnight_learner import _get_data_root
        import json as _json
        briefing_dir = _get_data_root() / ".edith" / "briefings"
        if not briefing_dir.exists():
            return {"sections": [], "message": "No briefings yet"}
        files = sorted(briefing_dir.glob("briefing_*.json"), reverse=True)
        if not files:
            return {"sections": [], "message": "No briefings yet"}
        with open(files[0]) as f:
            return _json.load(f)
    except Exception as e:
        return {"sections": [], "error": str(e)[:100]}


# ═══════════════════════════════════════════════════════════════════
# §9: SSE STREAMING — Live step-by-step execution for Agent Panel
# ═══════════════════════════════════════════════════════════════════

@router.post("/api/intent/stream", tags=["Agent"])
async def intent_stream(request: Request):
    """Stream intent classification + pipeline execution as SSE events.

    Events:
      - intent: classification result
      - step_start: beginning a pipeline step
      - step_done: step completed with results
      - done: all steps finished
    """
    from starlette.responses import StreamingResponse

    body = await request.json()
    query = body.get("query", body.get("text", ""))
    if not query:
        return _error(422, "no_query", "Provide a 'query'")

    async def generate():
        t0 = time.time()

        # Step 1: Classify intent
        from server.intent_router import route_intent
        intent_result = route_intent(query)
        pipeline_steps = intent_result.get("pipeline", [])

        yield f"event: intent\ndata: {_json_dumps_safe(intent_result)}\n\n"

        if not pipeline_steps:
            yield f"event: done\ndata: {_json_dumps_safe({'action': 'chat', 'total_elapsed_s': round(time.time() - t0, 2)})}\n\n"
            return

        # Step 2: Execute each step
        from server.routes.pipeline import _execute_step, _get_registry
        registry = _get_registry()
        context = {"text": query}

        for step_name in pipeline_steps:
            if step_name not in registry:
                continue

            yield f"event: step_start\ndata: {_json_dumps_safe({'step': step_name, 'label': registry.get(step_name, {}).get('label', step_name)})}\n\n"

            step_t0 = time.time()
            try:
                result = await _execute_with_retry(step_name, context, registry)
                elapsed = round(time.time() - step_t0, 2)

                for k, v in result.items():
                    if not k.startswith("_"):
                        context[k] = v

                retry_info = ""
                if result.get("self_corrected"):
                    retry_info = f" (self-corrected after {result.get('retries', 0)} retries)"

                yield f"event: step_done\ndata: {_json_dumps_safe({'step': step_name, 'elapsed_s': elapsed, 'result_keys': list(result.keys()), 'success': not result.get('error'), 'self_corrected': result.get('self_corrected', False), 'retries': result.get('retries', 0)})}\n\n"

            except Exception as e:
                elapsed = round(time.time() - step_t0, 2)
                yield f"event: step_done\ndata: {_json_dumps_safe({'step': step_name, 'elapsed_s': elapsed, 'error': str(e)[:200], 'success': False})}\n\n"

        total_elapsed = round(time.time() - t0, 2)
        yield f"event: done\ndata: {_json_dumps_safe({'total_elapsed_s': total_elapsed, 'steps_completed': len(pipeline_steps)})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


def _json_dumps_safe(obj: dict) -> str:
    """JSON-dump with fallback for non-serializable values."""
    import json as _json
    try:
        return _json.dumps(obj, default=str)
    except Exception:
        return _json.dumps({"error": "serialization_failed"})


# ═══════════════════════════════════════════════════════════════════
# §10: SESSION MEMORY — "What did we do last time?"
# ═══════════════════════════════════════════════════════════════════

@router.get("/api/agent/memory", tags=["Agent"])
async def get_session_memory():
    """Return recent topics, methods, and queries from research profile."""
    try:
        from server.research_profile import get_profile
        profile = get_profile()
        data = profile.get_profile()
        recent_queries = data.get("recent_queries", [])[-5:]
        top_topics = profile.get_top_topics(5)
        preferred_methods = profile.get_preferred_methods()[:5]
        tool_usage = sorted(
            data.get("tool_usage", {}).items(),
            key=lambda x: x[1], reverse=True
        )[:5]
        return {
            "recent_queries": recent_queries,
            "top_topics": top_topics,
            "preferred_methods": preferred_methods,
            "tool_usage": tool_usage,
            "context_summary": profile.get_context_string()[:500],
        }
    except Exception as e:
        return {"recent_queries": [], "top_topics": [], "error": str(e)[:100]}


# ═══════════════════════════════════════════════════════════════════
# §11: SELF-CORRECTION — Auto-retry failed steps with modified prompt
# ═══════════════════════════════════════════════════════════════════

async def _execute_with_retry(step_name: str, context: dict, registry: dict, max_retries: int = 2) -> dict:
    """Execute a pipeline step with self-correction on failure.

    If a step fails, modifies the context with error info and retries
    with a healing prompt — the model sees what went wrong and adjusts.
    """
    from server.routes.pipeline import _execute_step

    last_error = None
    for attempt in range(1 + max_retries):
        try:
            if attempt > 0 and last_error:
                context["_retry_attempt"] = attempt
                context["_previous_error"] = last_error
                context["_healing_instruction"] = (
                    f"Previous attempt failed with: {last_error}. "
                    f"Adjust your approach to avoid this error."
                )
                log.info(f"Self-correction: retrying '{step_name}' (attempt {attempt + 1})")

            result = await asyncio.to_thread(_execute_step, step_name, context)

            if result.get("error"):
                last_error = result["error"]
                if attempt < max_retries:
                    continue
                return {**result, "retries": attempt, "self_corrected": False}

            if attempt > 0:
                result["self_corrected"] = True
                result["retries"] = attempt
            return result

        except Exception as e:
            last_error = str(e)[:200]
            if attempt >= max_retries:
                return {"error": last_error, "retries": attempt, "self_corrected": False}

    return {"error": last_error or "max retries exceeded", "self_corrected": False}


# ═══════════════════════════════════════════════════════════════════
# §12: MULTI-AGENT DEBATE — Two personas argue a claim
# ═══════════════════════════════════════════════════════════════════

@router.post("/api/agent/debate", tags=["Agent"])
async def multi_agent_debate(request: Request):
    """Two AI personas debate different sides of a claim.

    Returns structured arguments from both sides + a synthesis verdict.
    """
    body = await request.json()
    claim = body.get("claim", body.get("query", ""))
    rounds = min(body.get("rounds", 2), 4)

    if not claim:
        return _error(422, "no_claim", "Provide a 'claim' to debate")

    from server.agentic import safe_llm_call

    advocate_system = (
        "You are the ADVOCATE — a political scientist who argues IN FAVOR of this claim. "
        "Be rigorous: cite mechanisms, provide evidence, use formal causal language. "
        "Respond concisely in 2-3 paragraphs."
    )
    critic_system = (
        "You are the CRITIC — a skeptical methodologist who argues AGAINST this claim. "
        "Challenge assumptions, identify confounders, demand better identification strategies. "
        "Respond concisely in 2-3 paragraphs."
    )

    debate_log = []
    t0 = time.time()

    for round_num in range(1, rounds + 1):
        # Advocate's turn
        prev_args = "\n".join(
            f"[{d['role']} Round {d['round']}]: {d['argument'][:300]}"
            for d in debate_log
        )
        advocate_prompt = (
            f"CLAIM: {claim}\n\n"
            f"{'PREVIOUS ARGUMENTS:\n' + prev_args + chr(10) if prev_args else ''}"
            f"Present your argument FOR this claim (Round {round_num})."
        )
        adv = safe_llm_call(advocate_prompt, system_instruction=advocate_system, temperature=0.7)
        debate_log.append({"role": "advocate", "round": round_num, "argument": adv["text"]})

        # Critic's turn
        prev_args = "\n".join(
            f"[{d['role']} Round {d['round']}]: {d['argument'][:300]}"
            for d in debate_log
        )
        critic_prompt = (
            f"CLAIM: {claim}\n\n"
            f"PREVIOUS ARGUMENTS:\n{prev_args}\n\n"
            f"Present your counter-argument AGAINST this claim (Round {round_num})."
        )
        critic = safe_llm_call(critic_prompt, system_instruction=critic_system, temperature=0.7)
        debate_log.append({"role": "critic", "round": round_num, "argument": critic["text"]})

    # Synthesis
    all_args = "\n\n".join(
        f"**{d['role'].upper()} (Round {d['round']}):**\n{d['argument']}"
        for d in debate_log
    )
    synthesis = safe_llm_call(
        f"CLAIM: {claim}\n\nDEBATE:\n{all_args}\n\n"
        "As a neutral judge, synthesize both sides. What is the strongest argument? "
        "What are the key unresolved questions? Give a verdict: SUPPORTED, CONTESTED, or UNSUPPORTED.",
        system_instruction="You are a neutral academic judge synthesizing a structured debate.",
        temperature=0.3,
    )

    return {
        "claim": claim,
        "rounds": rounds,
        "debate": debate_log,
        "synthesis": synthesis["text"],
        "elapsed_s": round(time.time() - t0, 2),
    }


# ═══════════════════════════════════════════════════════════════════
# §13: AGENT ARTIFACTS — Export pipeline results as Word/JSON
# ═══════════════════════════════════════════════════════════════════

@router.post("/api/agent/export", tags=["Agent"])
async def export_artifact(request: Request):
    """Export pipeline results as a downloadable report.

    Supported formats: 'json', 'word' (docx)
    """
    from starlette.responses import Response

    body = await request.json()
    data = body.get("data", {})
    title = body.get("title", "E.D.I.T.H. Agent Report")
    fmt = body.get("format", "json").lower()

    if fmt == "json":
        import json as _json
        content = _json.dumps(data, indent=2, default=str)
        return Response(
            content=content.encode(),
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{title}.json"'},
        )

    elif fmt == "word":
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(title, level=1)
            doc.add_paragraph(f"Generated by E.D.I.T.H. on {time.strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph("")

            # Add sections from data
            if isinstance(data, dict):
                for key, value in data.items():
                    if key.startswith("_"):
                        continue
                    doc.add_heading(key.replace("_", " ").title(), level=2)
                    if isinstance(value, list):
                        for item in value:
                            doc.add_paragraph(str(item), style="List Bullet")
                    elif isinstance(value, dict):
                        for k, v in value.items():
                            doc.add_paragraph(f"{k}: {v}")
                    else:
                        doc.add_paragraph(str(value))

            import io
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return Response(
                content=buf.read(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{title}.docx"'},
            )
        except ImportError:
            return _error(422, "no_docx", "python-docx not installed. Use format='json'.")

    return _error(400, "bad_format", f"Unknown format: {fmt}. Use 'json' or 'word'.")


# ═══════════════════════════════════════════════════════════════════
# §14: PERSISTENT QUEUE — Survives server restarts
# ═══════════════════════════════════════════════════════════════════

_QUEUE_FILE = None


def _get_queue_path():
    """Get persistent queue file path."""
    global _QUEUE_FILE
    if _QUEUE_FILE is None:
        try:
            from server.overnight_learner import _get_data_root
            queue_dir = _get_data_root() / ".edith"
            queue_dir.mkdir(parents=True, exist_ok=True)
            _QUEUE_FILE = queue_dir / "agent_queue.json"
        except Exception:
            from pathlib import Path
            _QUEUE_FILE = Path(".edith") / "agent_queue.json"
            _QUEUE_FILE.parent.mkdir(parents=True, exist_ok=True)
    return _QUEUE_FILE


def _load_queue() -> list[dict]:
    """Load queue from disk."""
    path = _get_queue_path()
    if path.exists():
        try:
            import json as _json
            with open(path) as f:
                return _json.load(f)
        except Exception:
            pass
    return []


def _save_queue(tasks: list[dict]) -> None:
    """Save queue to disk atomically."""
    import json as _json
    path = _get_queue_path()
    tmp = path.with_suffix(".tmp")
    try:
        with open(tmp, "w") as f:
            _json.dump(tasks[:50], f, indent=2, default=str)
        tmp.replace(path)
    except Exception as e:
        log.warning(f"Failed to save queue: {e}")


# Patch the existing queue endpoints to use persistent storage
_original_task_queue = _task_queue  # keep reference

def _init_persistent_queue():
    """Load persistent queue on startup."""
    global _task_queue
    loaded = _load_queue()
    if loaded:
        _task_queue.clear()
        _task_queue.extend(loaded)
        log.info(f"Loaded {len(loaded)} tasks from persistent queue")

# Auto-init on import
try:
    _init_persistent_queue()
except Exception:
    pass


def register(app, ns=None):
    """Register agent routes."""
    return router
